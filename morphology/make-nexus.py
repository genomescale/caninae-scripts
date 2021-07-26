import csv
import collections

from docx import Document

# The characters in the data set are ordered as follows (one-based):

#   1 to  11: Brain
#  12 to  87: Skull
#  88 to 192: Dentition
# 193 to 241: Body proportions and postcranial skeleton 
# 242 to 274: External morphology/soft anatomy
# 275 to 289: Metabolism, reproduction, development
# 290 to 346: Behaviour and ecology
# 347 to 360: Cytogenetics

# So essentially characters 12 through 241 are relevant to extant and extinct
# taxa, whereas other characters are only relevant to extant taxa. Of course if
# soft tissue fossils exist and have been studied then they could be included,
# but these characters are mostly absent from extinct taxa.

first_character_i = 12
last_character_i = 241

# one-based, fully closed range
def trim_sequence(sequence, start, end):
	trimmed_sequence = ""

	outside_brackets = True
	i = 1
	for c in sequence:
		if i >= start and i <= end:
			trimmed_sequence += c

		if c == "[":
			outside_brackets = False
		elif c == "]":
			outside_brackets = True

		if outside_brackets:
			i += 1

	return trimmed_sequence

label_mapping_filename = "label-mapping.csv"
label_mapping_file = open(label_mapping_filename)
label_mapping_reader = csv.reader(label_mapping_file)
label_mapping = {}
for label, species in label_mapping_reader:
	label_mapping[label] = species

state_keys_docx_filename = "zsc12293-sup-0021-appendixs1.docx"
state_values_docx_filename = "zsc12293-sup-0022-appendixs2.docx"

state_keys_docx_file = open(state_keys_docx_filename, "rb")
state_values_docx_file = open(state_values_docx_filename, "rb")

state_keys_docx = Document(state_keys_docx_file)
state_values_docx = Document(state_values_docx_file)

state_keys_docx_file.close()
state_values_docx_file.close()

character_descriptions = {}
character_keys = collections.defaultdict(list)

for paragraph in state_keys_docx.paragraphs:
	paragraph_text = paragraph.text.strip()
	paragraph_split = paragraph_text.split(None, 1)

	if len(paragraph_split) > 1:
		first_word = paragraph_split[0].rstrip(".")

		if first_word.isdigit():
			original_index = int(first_word)
			assert original_index not in character_descriptions

			description, keys = paragraph_text.split(None, 1)[1].split(":", 1)

			if "(" in description:
				character_descriptions[original_index] = description[:description.find("(")].strip()
			else:
				character_descriptions[original_index] = description.strip()

			key_breaks = []

			open_bracket = keys.find("(")
			while open_bracket != -1:
				close_bracket = keys.find(")", open_bracket + 1)
				if keys[open_bracket + 1:close_bracket].isdigit():
					key_breaks.append(close_bracket + 1)

				open_bracket = keys.find("(", open_bracket + 1)

			key_breaks.append(len(keys) - 1)

			for i in range(len(key_breaks) - 1):
				start = key_breaks[i]
				end = key_breaks[i + 1]

				state_key = keys[start:end].strip()
				if "(" in state_key:
					state_key = state_key[:state_key.rfind("(")].strip().rstrip(";")

				character_keys[original_index].append(state_key)

morphology = {}

for paragraph in state_values_docx.paragraphs:
	paragraph_split = paragraph.text.strip().split()
	if len(paragraph_split) == 2: # might be a sequence
		label, sequence = paragraph_split
		if label in label_mapping:
			binomial_name = label_mapping[label]
			truncated_sequence = trim_sequence(sequence, first_character_i, last_character_i).replace("[", "{").replace("]", "}")
			morphology[binomial_name] = truncated_sequence

nexus_template = """#NEXUS

BEGIN TAXA;
    DIMENSIONS NTAX=78;
    TAXLABELS
%s;
END;

BEGIN CHARACTERS;
    DIMENSIONS NCHAR=230;
    FORMAT DATATYPE=Standard SYMBOLS="0123" MISSING=? GAP=- INTERLEAVE=YES;
    CHARSTATELABELS
%s;
    MATRIX
%s;
END;
"""

taxa_block = ""
data_matrix = ""

for taxon in sorted(morphology):
	taxa_block += " " * 8 + taxon + "\n"
	data_matrix += " " * 8 + taxon + " " * (24 - len(taxon)) + "  " + morphology[taxon] + "\n"

characters_block = ""
for i in range(first_character_i, last_character_i + 1):
	characters_block += " " * 8 + str(1 + i - first_character_i) + " '" + character_descriptions[i] + "' /"
	for key in character_keys[i]:
		characters_block += " '" + key + "'"
	characters_block += ",\n"

nexus = nexus_template % (taxa_block, characters_block, data_matrix)
nexus_filename = "morphology.nex"
nexus_file = open(nexus_filename, "w")
nexus_file.write(nexus)
nexus_file.close()
